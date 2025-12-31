from flask import Flask, render_template, request, url_for, redirect
import random
import pandas as pd

app = Flask(__name__)


def convert_to_persian_digits(text):
    english_to_persian = str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹")
    return str(text).translate(english_to_persian)

def calculate_ders_scores(national_id, timestamp=None):

    # --- 1) Read Google Sheet
    sheet_id = "1vQ3BypmokEj06xEdLmPTSYYvExrVViIiUXn-ZfiJEDU"
    gid = "289405771"
    csv_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv&gid={gid}"
    df = pd.read_csv(csv_url, dtype=str, na_filter=False, keep_default_na=False)

    # --- 2) meta & items
    META_COLS = 6
    N_ITEMS   = 36
    reverse_items = {1,2,6,7,8,10,17,20,22,24,34}

    # --- 3) domains (1-based item numbers) — اصلاحش کن اگر با فرم فارسی‌ات فرق دارد
    domains = {
        'عدم پذیرش پاسخهای هیجانی (عدم پذیرش)': [11,12,21,23,25,29],
        'دشواری دست زدن به رفتار هدفمند (اهداف)': [13,18,20,26,33],
        'دشواری‌ کنترل تکانه (تکانه)':         [3,14,19,24,27,32],
        'فقدان آگاهی هیجانی (آگاهی)':             [2,6,8,10,17,34],
        'محدودیت در راهبردهای تنظیم هیجان (راهبردها)': [15,16,22,28,30,31,35,36],
        'فقدان شفافیت هیجانی (شفافیت)':          [1,4,5,7,9]
    }

    # helper: پاک‌سازی نیم‌فاصله از کلیدهای خروجی
    def _clean(s: str) -> str:
        return s.replace('\u200c', '').replace('‌','')  # ZWNJ & ZWJ

    # --- 4) locate participant row
    _norm_id = df['کد ملی'].astype(str).str.strip().str.replace('.0', '', regex=False)
    mask = _norm_id == str(national_id).strip()
    if timestamp is not None and 'Timestamp' in df.columns:
        mask = mask & (df['Timestamp'].astype(str).str.strip() == str(timestamp).strip())
    if not mask.any():
        raise ValueError("ردیفی با این کد ملی/تایم‌استمپ پیدا نشد.")
    row = df.loc[mask].iloc[0]
    rix = row.name

    # --- 5) numeric items 1..5 (بدها → 3)
    df_num = df.copy()
    item_cols = [df.columns[META_COLS + (i-1)] for i in range(1, N_ITEMS+1)]
    for c in item_cols:
        df_num[c] = pd.to_numeric(df[c], errors='coerce').clip(1,5).fillna(3).astype(int)

    # --- 6) reverse for this row
    for i, col in enumerate(item_cols, start=1):
        if i in reverse_items:
            df_num.at[rix, col] = int(6 - df_num.at[rix, col])

    # --- 7) compute per-domain & total
    domain_scores = {}
    for dname, items in domains.items():
        cols = [item_cols[i-1] for i in items]
        domain_scores[dname] = int(df_num.loc[rix, cols].sum())
    total_score = int(df_num.loc[rix, item_cols].sum())

    # --- 8) 3-band per domain (lower is better)
    def three_band(min_sum, max_sum, score):
        rng = max_sum - min_sum
        c1 = min_sum + 0.25 * rng
        c2 = min_sum + 0.75 * rng
        if score <= c1:   return 'خیر'   # خوب/پایین
        elif score <= c2: return 'سوسو'  # متوسط
        else:             return 'بله'   # بالا/نگران‌کننده

    # --- 9) outputs
    result = {
        'نام': row.get('نام و نام خانوادگی', ''),
        'جنسیت': row.get('جنسیت', ''),
        'نمره کل DERS': total_score
    }
    graph = {}  # فقط ۶ دامنه

    for dname, s in domain_scores.items():
        clean = _clean(dname)
        min_sum = 1 * len(domains[dname])
        max_sum = 5 * len(domains[dname])
        state = three_band(min_sum, max_sum, s)

        # در result: نمره + وضعیت هر دامنه
        result[f'نمره {clean}']   = s
        result[f'وضعیت {clean}']  = state
        # در graph: فقط نمره دامنه‌ها
        graph[f'نمره {clean}']    = s

    return [result, graph]


def classify_three_band(scale_name, score):
   
    # مثبت‌ها: بالاتر=بهتر | مسئله‌محورها: بالاتر=بدتر
    positive_scales = {'رفتار اجتماعی عملی', 'رفتار اجتماعی ارتباطی'}
    problem_scales  = {'رفتار ضد اجتماعی آشکار', 'رفتار ضد اجتماعی رابطه ای', 'قربانی'}

    if scale_name in positive_scales:
        # 0–1 پایین، 2–5 هنجار (ادغام دوم و سوم)، 6–8 بالاتر از هنجار
        if score <= 1:
            return 'بله'
        elif 2 <= score <= 5:
            return 'سوسو'
        else:  # 6–8
            return 'خیر'

    if scale_name in problem_scales:
        # 0–1 پایین‌تر از هنجار (مشکل کم/خوب)، 2–5 هنجار، 6–8 بالاتر از هنجار (نگران‌کننده)
        if score <= 1:
            return 'خیر'
        elif 2 <= score <= 5:
            return 'سوسو'
        else:  # 6–8
            return 'بله'

    # پیش‌فرض محافظه‌کارانه
    return 'هنجار'


import pandas as pd

def calculate_csbs_scores(national_id, timestamp):
    # 1) Load sheet
    sheet_id = "1rK_u0GUij4k7IGxJcUTAvJ_PhIzJsngZWBk59xbdakM"
    gid = "384146557"
    csv_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv&gid={gid}"
    df = pd.read_csv(csv_url, dtype=str, na_filter=False, keep_default_na=False)

    # 2) Score map
    score_map = {'هرگز': 0, 'بعضی اوقات': 1, 'اغلب اوقات': 2}

    # 3) Subscales (question numbers)
    scales = {
        'رفتار اجتماعی عملی':            [1, 8, 11, 16],
        'رفتار اجتماعی ارتباطی':        [5, 14, 18, 22],
        'رفتار ضد اجتماعی آشکار':       [4, 7, 12, 20],
        'رفتار ضد اجتماعی رابطه ای':    [17, 19, 23, 24],
        'قربانی':                        [3, 6, 9, 13],
    }
    # columns are positional: add 4 because first 5 columns are metadata (0..4)
    scales_index = {k: [q + 4 for q in v] for k, v in scales.items()}

    # helper: clean key (remove ZWNJ for display keys only)
    def _clean_key(s: str) -> str:
        return s.replace('\u200c', ' ')

    # 4) Filter rows directly (no masks/index headaches)
    df = df[
        (df['کد ملی'].astype(str).str.strip().str.replace('.0', '', regex=False) == str(national_id).strip()) &
        (df['Timestamp'].astype(str).str.strip() == str(timestamp).strip())
    ]
    if df.empty:
        raise ValueError("ردیفی با این کد ملی/تایم‌استمپ پیدا نشد.")

    # We'll compute scores from this single-row slice
    row = df.iloc[0]     # the matched row
    rpos = 0             # row position in this filtered df

    # 5) Map string answers to numbers ONLY on the filtered frame
    df_numeric = df.copy()
    for col in df_numeric.columns[5:]:
        df_numeric[col] = df_numeric[col].map(score_map).fillna(0).astype(int)

    # 6) Build scores
    category_scores = {}
    for scale_name, idxs in scales_index.items():
        score_val = int(df_numeric.iloc[rpos, idxs].sum())
        clean_scale = _clean_key(scale_name)
        category_scores['نمره ' + clean_scale] = score_val

    # 7) Final results (metadata + scores + status)
    results = {
        'نام': row.get('نام و نام خانوادگی', ''),
        'جنسیت': row.get('جنسیت', ''),
    }
    for scale_name, _ in scales_index.items():
        clean_scale = _clean_key(scale_name)
        score_val = category_scores['نمره ' + clean_scale]
        results['نمره ' + clean_scale] = score_val
        # برای کلاس‌بندی از نام اصلی (ممکن است ZWNJ داشته باشد) استفاده کن
        results['وضعیت ' + clean_scale] = classify_three_band(scale_name, score_val)

    return results, category_scores


def calculate_cbcl_scores(national_id):
    return {"عامل ۱": random.randint(20, 60), "عامل ۲": random.randint(20, 60), "عامل ۳": random.randint(20, 60)}


def calculate_scrs_scores(national_id, timestamp):
    csv_url = 'https://docs.google.com/spreadsheets/d/17QsSqNPQ2a1ZFgYmwAH1UBS8btZe7RAHTxFcAZMHe6A/export?format=csv&gid=1991538008'
    df = pd.read_csv(csv_url, dtype=str, na_filter=False, keep_default_na=False)
    df = df[
    (df['کد ملی'].astype(str).str.strip().str.replace('.0', '', regex=False) == national_id) &
    (df['Timestamp'] == timestamp) 
    ]
    if not df[df['کد ملی'].astype(str).str.strip().str.replace('.0', '', regex=False) == national_id]['جنسیت'].values[0]:
        return {"error": f"دانش‌آموزی با کد ملی '{national_id}' یافت نشد."}
    ranges = [
    {'range': [33, 58], 'color': 'lightgreen', 'label': 'خیر'},
    {'range': [59, 142], 'color': 'yellow', 'label': 'سوسو'},
    {'range': [143, 231], 'color': 'red', 'label': 'بله'}
]

    result_label = None
    for r in ranges:
        low, high = r['range']
        if low <= int(df['Total Score']) <= high:
            result_label = r['label']
            break
    result = {

        'نام': df[df['کد ملی'].astype(str).str.strip().str.replace('.0', '', regex=False) == national_id]['نام و نام خانوادگی'].values[0],
        'نمره خودکنترلی': int(df['Total Score']),
        'وضعیت خودکنترلی' : result_label
    }
    graph = {
        'نمره خودکنترلی': int(df['Total Score'])
    }
    return [result, graph]

def calculate_anger_scores(national_id, timestamp):
    csv_url = (
            "https://docs.google.com/spreadsheets/d/1bpZjmiBsEnriX6y34ryvtho3vhn7rQ9t13kWyQA0cVQ/"
            "export?format=csv&gid=420825805&cachebust=" + str(int(time.time()))
        )

        # فقط این پارامترها کافی‌ان تا «nan» و تبدیل عددی رخ نده
    df = pd.read_csv(csv_url, dtype=str, na_filter=False, keep_default_na=False)
        # Score map
    score_map = {
        'هرگز یا به ندرت': 1,
        'یک‌ بار در ماه': 2,
        'یک بار در هفته': 3,
        'اغلب روزا': 4,
    }
    
    # Identify columns for each factor
    columns = df.columns
    factor_1_cols = columns[10:17]  # Q7-Q13
    factor_2_cols = columns[17:26]  # Q14-Q21
    factor_3_cols = columns[4:10]   # Q1-Q6
    df = df[
    (df['کد ملی'].astype(str).str.strip().str.replace('.0', '', regex=False) == national_id) &
    (df['Timestamp'] == timestamp) 
    ]
    if not df[df['کد ملی'].astype(str).str.strip().str.replace('.0', '', regex=False) == national_id]['جنسیت'].values[0]:
        return {"error": f"دانش‌آموزی با کد ملی '{national_id}' یافت نشد."}
    if len(df) == 1:
        def score_factor(cols):
            return sum(score_map.get(str(df[col].values[0]).strip(), 0) for col in cols)

        f1 = score_factor(factor_1_cols)
        f2 = score_factor(factor_2_cols)
        f3 = score_factor(factor_3_cols)

        f1_threshold = 8 if df[df['کد ملی'].astype(str).str.strip().str.replace('.0', '', regex=False) == national_id]['جنسیت'].values[0] == 'دختر' else 10
        f2_threshold = 18 if df[df['کد ملی'].astype(str).str.strip().str.replace('.0', '', regex=False) == national_id]['جنسیت'].values[0] == 'دختر' else 17
        f3_threshold = 15 if df[df['کد ملی'].astype(str).str.strip().str.replace('.0', '', regex=False) == national_id]['جنسیت'].values[0] == 'دختر' else 16

        # تبدیل به بله/خیر
        def boolean_to_farsi(val):
            return 'بله' if val else 'خیر'

        result = {
            'نام': df[df['کد ملی'].astype(str).str.strip().str.replace('.0', '', regex=False) == national_id]['نام و نام خانوادگی'].values[0],
            'جنسیت': df[df['کد ملی'].astype(str).str.strip().str.replace('.0', '', regex=False) == national_id]['جنسیت'].values[0],
            'نمره پرخاش جسمانی': f1,
            'وضعیت پرخاش جسمانی': boolean_to_farsi(f1 > f1_threshold),
            'نمره پرخاش رابطه‌ای': f2,
            'وضعیت پرخاش رابطه‌ای': boolean_to_farsi(f2 > f2_threshold),
            'نمره واکنشی کلامی': f3,
            'وضعیت واکنشی کلامی': boolean_to_farsi(f3 > f3_threshold)
        }
        graph = {
            'نمره پرخاش جسمانی': f1,
            'نمره پرخاش رابطه‌ای': f2,
            'نمره واکنشی کلامی': f3,
        }
        return [result, graph]

    else:
        return {'error': 'this feature is not ready yet!'}


import plotly.graph_objects as go
import os
import plotly.graph_objects as go
import os
import plotly.graph_objects as go

def create_ders_gauge_chart(national_id, score, factor_index, factor):
    """
    همسان با create_csbs_gauge_chart:
      inputs: (national_id, score, factor_index, factor)
      output: 'national_id/gauge_{factor_index}.png'
    - دامنه‌های DERS: رنج گیج = تعداد آیتم * [1..5]
    - رنگ‌بندی مسئله‌محور: پایین=سبز، میانی=زرد، بالا=قرمز
    """

    # پوشه و نام فایل مثل CSBS
    folder_path = f'static/{national_id}'
    os.makedirs(folder_path, exist_ok=True)
    filename = f'{folder_path}/gauge_{factor_index}.png'
    rel = f'{national_id}/gauge_{factor_index}.png'

    # نگاشت دامنه‌ها به تعداد آیتم‌ها (کلیدها هم با و هم بی "نمره " پشتیبانی می‌شوند)
    domain_items = {
        'عدم پذیرش پاسخ‌های هیجانی (عدم پذیرش)': 6,
        'دشواری دست زدن به رفتار هدفمند (اهداف)': 5,
        'دشواری کنترل تکانه (تکانه)': 6,
        'فقدان آگاهی هیجانی (آگاهی)': 6,
        'محدودیت در راهبردهای تنظیم هیجان (راهبردها)': 8,
        'فقدان شفافیت هیجانی (شفافیت)': 5,
    }

    # اگر factor به صورت "نمره ..." آمده باشد، برچسب دامنه را استخراج کن
    fac_clean = str(factor).strip()
    if fac_clean.startswith('نمره '):
        fac_key = fac_clean.replace('نمره ', '', 1).strip()
    else:
        fac_key = fac_clean

    # تلاش برای یافتن دامنه؛ اگر پیدا نشد، پیش‌فرض 6 آیتم (رنج 6..30)
    n_items = domain_items.get(fac_key, None)
    if n_items is None:
        # تلاش تطبیق ساده بدون نیم‌فاصله
        def _simp(s): return s.replace('\u200c','').replace('‌','').strip()
        simp_key = _simp(fac_key)
        found = None
        for k in domain_items:
            if _simp(k) == simp_key:
                found = k
                break
        n_items = domain_items.get(found, 6)

    axis_min = 1 * n_items
    axis_max = 5 * n_items

    # تقسیم سه‌گانه 25%/75%
    span = axis_max - axis_min
    b1_hi = axis_min + 0.25 * span
    b2_hi = axis_min + 0.75 * span

    # مسئله‌محور: پایین=سبز، میانی=زرد، بالا=قرمز
    steps = [
        {'range': [axis_min, b1_hi], 'color': 'lightgreen'},
        {'range': [b1_hi,  b2_hi],  'color': 'yellow'},
        {'range': [b2_hi,  axis_max], 'color': 'red'}
    ]
    bar_color = "darkred"

    # کلَمپ و ساخت گیج
    try:
        val = float(score)
    except Exception:
        val = axis_min
    val = max(axis_min, min(axis_max, val))

    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=val,
        title={'text': f'{factor}', 'font': {'size': 20, 'family': 'Vazirmatn'}},
        gauge={
            'axis': {'range': [axis_min, axis_max], 'dtick': max(1, round(span/8))},
            'bar': {'color': bar_color},
            'steps': steps,
            'threshold': {'line': {'color': "black", 'width': 4}, 'value': val}
        }
    ))
    fig.write_image(filename)  # نیاز به kaleido
    return rel


def create_scrs_gauge_chart(national_id, score, factor_index):

    title = 'نمره خودکنترلی'
    folder_path = f'static/{national_id}'
    os.makedirs(folder_path, exist_ok=True)  # ✅ Create if it doesn't exist
    filename = f'static/{national_id}/gauge_scrs.png'
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=score,
        title={'text': title, 'font': {'size': 20, 'family': 'Vazirmatn'}},
        gauge={
            'axis': {'range': [33, 231]},
            'bar': {'color': "darkblue"},
            'steps': [
                {'range': [33, 58], 'color': 'lightgreen'},
                {'range': [59, 142], 'color': 'yellow'},
                {'range': [143, 231], 'color': 'red'}
            ],
            'threshold': {
                'line': {'color': "black", 'width': 4},
                'value': score
            }
        }
    ))
    fig.write_image(filename)
    return f'{national_id}/gauge_scrs.png'


def create_anger_gauge_chart(national_id, score, gender, factor_index, factor):

    title = factor
    folder_path = f'static/{national_id}'
    os.makedirs(folder_path, exist_ok=True)  # ✅ Create if it doesn't exist
    filename = f'static/{national_id}/gauge_anger_{factor_index}.png'

    if factor_index == 1:
        if gender == 'پسر':
            steps = [
                {'range': [0, 11], 'color': 'lightgreen'},
                {'range': [11, 22], 'color': 'red'}
            ]
        else:
            steps = [
                {'range': [0, 9], 'color': 'lightgreen'},
                {'range': [9, 22], 'color': 'red'}
            ]
        range = [0, 22]
    elif factor_index == 2:
        if gender == 'پسر':
            steps = [
                {'range': [0, 18], 'color': 'lightgreen'},
                {'range': [18, 25], 'color': 'red'}
            ]
        else:
            steps = [
                {'range': [0, 19], 'color': 'lightgreen'},
                {'range': [19, 25], 'color': 'red'}
            ]
        range = [0, 25]

    else:
        if gender == 'پسر':
            steps = [
                {'range': [0, 16], 'color': 'lightgreen'},
                {'range': [16, 19], 'color': 'red'}
            ]
        else:
            steps = [
                {'range': [0, 17], 'color': 'lightgreen'},
                {'range': [17, 19], 'color': 'red'}
            ]
        range = [0, 19]

    
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=score,
        title={'text': title, 'font': {'size': 20, 'family': 'Vazirmatn'}},
        gauge={
            'axis': {'range': range},
            'bar': {'color': "darkblue"},
            'steps': steps,
            'threshold': {
                'line': {'color': "black", 'width': 4},
                'value': score
            }
        }
    ))
    fig.write_image(filename)
    return f'{national_id}/gauge_anger_{factor_index}.png'



def create_csbs_gauge_chart(national_id, score, factor_index, factor):
    # Create folder
    folder_path = f'static/{national_id}'
    os.makedirs(folder_path, exist_ok=True)

    # File name
    filename = f'{folder_path}/gauge_{factor_index}.png'

    # Decide ranges depending on scale type
    if factor_index in [1, 2]:
        # Positive scales: low=red, mid=yellow, high=green
        steps = [
            {'range': [0, 2], 'color': 'red'},         # 0–1
            {'range': [2, 6], 'color': 'yellow'},      # 2–5
            {'range': [6, 8], 'color': 'lightgreen'}   # 6–8
        ]
        bar_color = "darkblue"
    else:
        # Problem scales: low=green, mid=yellow, high=red
        steps = [
            {'range': [0, 2], 'color': 'lightgreen'},  # 0–1
            {'range': [2, 6], 'color': 'yellow'},      # 2–5
            {'range': [6, 8], 'color': 'red'}          # 6–8
        ]
        bar_color = "darkred"

    # Build gauge
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=score,
        title={'text': f' {factor}', 'font': {'size': 20, 'family': 'Vazirmatn'}},
        gauge={
            'axis': {'range': [0, 8], 'dtick': 1},
            'bar': {'color': bar_color},
            'steps': steps,
            'threshold': {
                'line': {'color': "black", 'width': 4},
                'value': score
            }
        }
    ))
    # Save image
    fig.write_image(filename)
    return f'{national_id}/gauge_{factor_index}.png'


# Mapping of test codes to function references
test_functions = {
    "DERS": calculate_ders_scores,
    "CSBS": calculate_csbs_scores,
    "CBCL": calculate_cbcl_scores,
    "SCRS": calculate_scrs_scores,
    "ANGER": calculate_anger_scores
}

# Display names (Persian)
test_names = {
    "DERS": "ابزار سنجش هوش هیجانی",
    "CSBS": "ابزار سنجش هوش اجتماعی",
    "CBCL": "چک‌لیست رفتاری کودکان",
    "SCRS": "ابزار سنجش قانون پذیری و تعهد",
    "ANGER": "ابزار سنجش مهارت های مدیریت پرخاش"
}

# ----------------- Routes --------------------
@app.route("/results_page", methods=["GET"])
def results_page():
    print("Serving index page--------------------------")
    return render_template("index.html")

@app.route("/", methods=["GET"])
def home():
    print("Serving index page--------------------------")
    return render_template("first_page.html")

from docx import Document
from docx.shared import Inches
import os

import os
from docx import Document
from docx.shared import Inches


import time
import re

def clean_digits(s: str) -> str:
    s = str(s or "")
    # حذف آپاستروف ابتدای مقدار (متن اجباری شیتس)
    if s.startswith("'"):
        s = s[1:]
    # حذف کاراکترهای جهت‌دهی و فاصله‌های نامرئی
    s = (s.replace("\u200e","")  # LRM
         .replace("\u200f","")   # RLM
         .replace("\u202a","").replace("\u202b","").replace("\u202c","")
         .replace("\u202d","").replace("\u202e","")
         .replace("\u00a0"," ")  # NBSP
         .strip())
    # فقط رقم‌ها را نگه دار (اگر می‌خواهی + هم بماند، الگو را عوض کن)
    return re.sub(r"\D", "", s)

@app.route('/choose/<test_name>', methods=['POST'])
def choose(test_name):
    national_id = request.form.get('national_id')

    if test_name == 'ANGER':
       csv_url = (
            "https://docs.google.com/spreadsheets/d/1bpZjmiBsEnriX6y34ryvtho3vhn7rQ9t13kWyQA0cVQ/"
            "export?format=csv&gid=420825805&cachebust=" + str(int(time.time()))
        )

        # فقط این پارامترها کافی‌ان تا «nan» و تبدیل عددی رخ نده
       df = pd.read_csv(csv_url, dtype=str, na_filter=False, keep_default_na=False, encoding="utf-8")
       row = df[df['کد ملی'].astype(str).str.strip().str.replace('.0', '', regex=False) == national_id]

       row = row.to_dict(orient="records")
       return render_template(
        "view_attempts.html",
        test_code="ANGER",
        national_id=national_id,
        attempts=row,
        test_names=test_names
    )
    #render_template()
    elif test_name == 'SCRS':
       csv_url = 'https://docs.google.com/spreadsheets/d/17QsSqNPQ2a1ZFgYmwAH1UBS8btZe7RAHTxFcAZMHe6A/export?format=csv&gid=1991538008'
       df = pd.read_csv(csv_url, dtype=str, na_filter=False, keep_default_na=False)
       row = df[df['کد ملی'].astype(str).str.strip().str.replace('.0', '', regex=False) == national_id]
       row = row.to_dict(orient="records")
       return render_template(
        "view_attempts.html",
        test_code="SCRS",
        national_id=national_id,
        attempts=row,
        test_names=test_names
        
    )
    elif test_name == 'CBCL':
        # Handle CBCL test
        return f"Processing CBCL test for {national_id}"
    
    elif test_name == 'CSBS':
       sheet_id = "1rK_u0GUij4k7IGxJcUTAvJ_PhIzJsngZWBk59xbdakM"
       gid = "384146557"
       csv_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv&gid={gid}"
       df = pd.read_csv(csv_url, dtype=str, na_filter=False, keep_default_na=False)
       row = df[df['کد ملی'].astype(str).str.strip().str.replace('.0', '', regex=False) == national_id]
       row = row.to_dict(orient="records")
       return render_template(
        "view_attempts.html",
        test_code="CSBS",
        national_id=national_id,
        attempts=row,
        test_names=test_names
        )
        
    elif test_name == 'DERS':
       sheet_id = "1vQ3BypmokEj06xEdLmPTSYYvExrVViIiUXn-ZfiJEDU"
       gid = "289405771"
       csv_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv&gid={gid}"
       df = pd.read_csv(csv_url, dtype=str, na_filter=False, keep_default_na=False)
       row = df[df['کد ملی'].astype(str).str.strip().str.replace('.0', '', regex=False) == national_id]
       row = row.to_dict(orient="records")
       return render_template(
        "view_attempts.html",
        test_code="DERS",
        national_id=national_id,
        attempts=row,
        test_names=test_names
        )
    else:
        return "Invalid test name", 400

import json
import jdatetime
from datetime import datetime

# تبدیل اعداد انگلیسی به فارسی

PERSIAN_MONTHS = {
    1: "فروردین",
    2: "اردیبهشت",
    3: "خرداد",
    4: "تیر",
    5: "مرداد",
    6: "شهریور",
    7: "مهر",
    8: "آبان",
    9: "آذر",
    10: "دی",
    11: "بهمن",
    12: "اسفند"
}

# تبدیل اعداد انگلیسی به فارسی
def convert_to_persian_digits(text):
    english_to_persian = str.maketrans("0123456789", "۰۱۲۳۴۵۶۷۸۹")
    return str(text).translate(english_to_persian)

# تابع نهایی برای تبدیل تاریخ-ساعت میلادی به تاریخ شمسی فارسی
def datetime_to_persian_date(datetime_str):
    # تبدیل رشته به datetime
    dt = datetime.strptime(datetime_str, "%d/%m/%Y %H:%M:%S")

    # تبدیل تاریخ میلادی به شمسی
    jalali_date = jdatetime.date.fromgregorian(date=dt.date())

    # گرفتن روز، ماه، سال شمسی
    day_persian = convert_to_persian_digits(jalali_date.day)
    month_persian = PERSIAN_MONTHS[jalali_date.month]
    year_persian = convert_to_persian_digits(jalali_date.year)

    return f"{day_persian} {month_persian} {year_persian}"

@app.route("/results", methods=["POST"])
def results():
    test_code = request.form.get("test_code")
    national_id = request.form.get("national_id")
    timestamp = request.form.get("timestamp")
    age = request.form.get("age")
    with open("test_descriptions.json", "r", encoding="utf-8") as f:
        DESCRIPTIONS = json.load(f)
    with open("test_def.json", "r", encoding="utf-8") as f:
        DEF = json.load(f)
    if test_code not in test_functions:
        return render_template("index.html", tests=test_names, result={"error": "آزمون نامعتبر است"})
    graph = None
    calculate_fn = test_functions[test_code]
    scores = calculate_fn(national_id, timestamp)
    date = datetime_to_persian_date(request.form.get('timestamp'))

    if test_code == 'ANGER':
        scores, graph = scores
        chart_filenames = {}
        gender = request.form.get("gender")
        phone = request.form.get("phone")
        DESCRIPTIONS = DESCRIPTIONS['tests']['ANGER']
        for i, factor in enumerate([ 'نمره پرخاش جسمانی',  'نمره پرخاش رابطه‌ای',  'نمره واکنشی کلامی'], start=1):
            chart_filename = create_anger_gauge_chart(national_id, scores[factor], gender,i, factor)
            chart_filenames[factor] = chart_filename

    elif test_code == 'SCRS':
        scores, graph = scores
        chart_filenames = {}
        DESCRIPTIONS = DESCRIPTIONS['tests']['SCRS']
        phone = request.form.get("phone")
        for i, factor in enumerate(["نمره خودکنترلی"], start=1):
            chart_filename = create_scrs_gauge_chart(national_id, scores[factor], i)
            chart_filenames[factor] = chart_filename

    if test_code == 'CSBS':
        scores, graph = scores
        chart_filenames = {}
        phone = request.form.get("phone")
        DESCRIPTIONS = DESCRIPTIONS['tests']['CSBS']
        for i, factor in enumerate(['نمره رفتار اجتماعی عملی', 'نمره رفتار اجتماعی ارتباطی', 'نمره رفتار ضد اجتماعی آشکار', 'نمره رفتار ضد اجتماعی رابطه ای', 'نمره قربانی'], start=1):
            factor = factor.replace('\u200c', ' ')
            chart_filename = create_csbs_gauge_chart(national_id, scores[factor], i, factor)
            chart_filenames[factor] = chart_filename
    if test_code == 'DERS':
        scores, graph = scores
        chart_filenames = {}
        phone = request.form.get("phone")
        DESCRIPTIONS = DESCRIPTIONS['tests']['DERS']
        for i, factor in enumerate(['نمره عدم پذیرش پاسخهای هیجانی (عدم پذیرش)', 'نمره دشواری دست زدن به رفتار هدفمند (اهداف)', 'نمره دشواری کنترل تکانه (تکانه)', 'نمره فقدان آگاهی هیجانی (آگاهی)', 'نمره محدودیت در راهبردهای تنظیم هیجان (راهبردها)', 'نمره فقدان شفافیت هیجانی (شفافیت)'], start=1):           
            factor = factor.replace('\u200c', ' ')
            chart_filename = create_ders_gauge_chart(national_id, scores[factor], i, factor)
            chart_filenames[factor] = chart_filename
    
    return render_template(
        "results.html",
        name=scores['نام'],
        test_code=test_code,
        test_name=test_names[test_code],
        national_id=national_id,
        age=age,
        phone=phone,
        timestamp=date,
        graph=graph,
        result=scores,
        charts=chart_filenames,
        descriptions=DESCRIPTIONS,
        DEF=DEF,

    )


from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.shape import WD_INLINE_SHAPE
from flask import request, render_template, send_file
from io import BytesIO
from docx import Document
from zipfile import ZipFile, BadZipFile
from docx.shared import Pt

from htmldocx import HtmlToDocx


def tidy_text(text):
    """اصلاح فاصله‌ها و جلوگیری از چسبیدن کلمات فارسی"""
    text = text.replace(" :", " : ").replace("  ", " ")
    text = text.replace("نمره:", "نمره :")
    text = text.replace("وضعیت:", "وضعیت :")
    return text.strip()

def set_paragraph_rtl(paragraph):
    """تنظیم راست‌چین و راست به چپ برای پاراگراف"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()

    # bidi راست به چپ
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

    # تراز به راست
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    pPr.append(jc)

def split_score_and_status(doc):
    """جدا کردن نمره و وضعیت در پاراگراف‌های جدا (نسخه بدون ارور)"""
    paragraphs_copy = list(doc.paragraphs)  # کپی امن از لیست پاراگراف‌ها
    for p in paragraphs_copy:
        text = p.text.strip()
        if "وضعیت" in text and "نمره" in text:
            # جدا کردن بخش نمره و وضعیت
            parts = text.split("وضعیت", 1)
            score_part = parts[0].strip()
            status_part = "وضعیت " + parts[1].strip()

            # پاراگراف اول (نمره)
            p.clear()
            p.add_run(score_part)
            set_paragraph_rtl(p)

            # پاراگراف دوم (وضعیت) → درست بعد از p اضافه می‌کنیم
            para = p.insert_paragraph_before(status_part)
            set_paragraph_rtl(para)


def fix_spacing_and_order(doc):
    for p in doc.paragraphs:
        p.paragraph_format.space_before = Pt(0)  # بدون فاصله قبل
        p.paragraph_format.space_after = Pt(0)   # بدون فاصله بعد
        text = p.text.strip()
        if "نمره" in text and "وضعیت" not in text:
            # شکستن متن به دو خط: عنوان و مقدار نمره
            parts = text.split("نمره")
            if len(parts) == 2:
                title = parts[0].strip()
                score = parts[1].strip(" :")
                
                # پاراگراف عنوان
                p.text = title + "    نمره: " + score
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in p.runs:
                    run.font.size = Pt(12)
                    
        if "وضعیت" in text:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def style_document(doc):
    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        # تنظیم فونت پیش‌فرض و اندازه
        for run in paragraph.runs:
            run.font.name = 'Vazirmatn'  # یا مثلاً 'Arial'
            run.font.size = Pt(12)
            # برای بولد کردن:
            run.bold = True
        
        # راست‌چین کردن همه پاراگراف‌ها
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

@app.route('/results-docx', methods=['POST'])
def results_docx():
    images = "header.png"
    
    data = request.get_json(force=True)
    with open("test_descriptions.json", "r", encoding="utf-8") as f:
        DESCRIPTIONS = json.load(f)
        DESCRIPTIONS = DESCRIPTIONS['tests'][data.get('test_code')]
    with open("test_def.json", "r", encoding="utf-8") as f:
        DEF = json.load(f)
    html = render_template(
        'results.html',
        name=data.get('name'),
        test_name=data.get('test_name'),
        national_id=data.get('national_id'),
        phone=data.get('phone'),
        result=data.get('result'),
        timestamp=data.get('timestamp'),
        age=data.get('age'),
        graph=data.get('graph'),
        charts=data.get('charts'),
        imgs=images,
        descriptions=DESCRIPTIONS,
        DEF=DEF,
        for_docx=True
    )

    # ایجاد فایل Word
    doc = Document()
    converter = HtmlToDocx()

    # افزودن جهت RTL به HTML
    full_html = (
        '<!DOCTYPE html><html lang="fa" dir="rtl">'
        '<head><meta charset="utf-8"></head>'
        '<body>' + html + '</body></html>'
    )
    converter.add_html_to_document(full_html, doc)
    style_document(doc)
    fix_spacing_and_order(doc)
    # اصلاح متن و راست‌چین کردن
    for p in doc.paragraphs:
        fixed = tidy_text(p.text)
        if fixed != p.text:
            p.clear()
            p.add_run(fixed)
        set_paragraph_rtl(p)

    # جدا کردن نمره و وضعیت
    split_score_and_status(doc)

    # ذخیره در حافظه
    buf = BytesIO()
    doc.save(buf)

    # بررسی صحت فایل DOCX
    buf.seek(0)
    try:
        ZipFile(buf).testzip()
    except BadZipFile:
        raise RuntimeError("Generated DOCX is invalid ZIP.")
    buf.seek(0)

    TARGET_WIDTH = Inches(2.4)  # ~12.2 cm — pick what you like
    for ishape in doc.inline_shapes:
        if ishape.type == WD_INLINE_SHAPE.PICTURE:
            ishape.width = TARGET_WIDTH 
    filename = f"نتایج-آزمون-{data.get('test_name','')}-{data.get('name','')}-{data.get('national_id','')}.docx"

    folder_path =  'static/' + data.get('national_id')  # مسیر پوشه‌ای که می‌خوای پاک شه

    if os.path.exists(folder_path) and os.path.isdir(folder_path):
        os.system(f'rm -rf "{folder_path}"') 
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ----------------- Run the App --------------------
if __name__ == "__main__":
    app.run(debug=True)
